"""
Layer: INFRASTRUCTURE
Imports allowed: domain + pypdf2 + python-docx
Purpose: Document parsing and sliding-window chunking.
"""
import logging
from pathlib import Path

from domain.exceptions import DocumentParseError
from domain.interfaces import IChunker, IDocumentParser
from domain.models import Document, SensitivityLevel

logger = logging.getLogger(__name__)


class DocumentParser(IDocumentParser):
    """Parses PDF, DOCX, and TXT files into Document objects."""

    def parse(self, filepath: str) -> Document:
        path = Path(filepath)
        if not path.exists():
            raise DocumentParseError(f"File not found: {filepath}")

        suffix = path.suffix.lower()
        try:
            if suffix == ".pdf":
                content = self._parse_pdf(path)
            elif suffix == ".docx":
                content = self._parse_docx(path)
            elif suffix in (".txt", ".md"):
                content = path.read_text(encoding="utf-8")
            else:
                raise DocumentParseError(f"Unsupported format: {suffix}")

            return Document(
                filename=path.name,
                content=content,
                sensitivity=SensitivityLevel.HIGH,  # default HIGH until classified
                encrypted=False,  # encryption applied by encryptor separately
            )
        except DocumentParseError:
            raise
        except Exception as e:
            raise DocumentParseError(f"Parse failed for {filepath}: {e}") from e

    @staticmethod
    def _parse_pdf(path: Path) -> str:
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            raise DocumentParseError("pypdf not installed")

    @staticmethod
    def _parse_docx(path: Path) -> str:
        try:
            import docx
            doc = docx.Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            raise DocumentParseError("python-docx not installed")


class SlidingWindowChunker(IChunker):
    """
    Sliding window chunker with overlap.
    Splits on word boundaries to avoid cutting mid-sentence.
    """

    def __init__(self, chunk_size: int = 512, overlap: int = 50) -> None:
        self._chunk_size = chunk_size
        self._overlap = overlap

    def chunk(self, document: Document) -> list[str]:
        words = document.content.split()
        if not words:
            return []

        chunks = []
        start = 0
        while start < len(words):
            end = min(start + self._chunk_size, len(words))
            chunk = " ".join(words[start:end])
            if chunk.strip():
                chunks.append(chunk)
            if end >= len(words):
                break
            start = end - self._overlap  # overlap for context continuity

        logger.debug("Chunked '%s': %d chunks", document.filename, len(chunks))
        return chunks
